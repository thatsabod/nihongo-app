from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE_TEXT = ROOT / "pdf_text_extract.txt"
OUTPUT = ROOT / "اليابانية من طوكيو - نسخة محسنة.docx"

LESSON = "\u0627\u0644\u062f\u0631\u0633"
TERMS = "\u0627\u0644\u0645\u0635\u0637\u0644\u062d\u0627\u062a"
PATTERNS = "\u0623\u0646\u0645\u0627\u0637  \u0627\u0644\u062c\u0645\u0644\u0629"
TRANSLATION = "\u0627\u0644\u062a\u0631\u062c\u0645\u0629 II"
CONVERSATION = "\u0627\u0644\u0645\u062d\u0627\u062f\u062b\u0629"
GRAMMAR = "\u0627\u0644\u0642\u0648\u0627\u0639\u062f"

COLORS = {
    "ink": "172033",
    "muted": "5F6678",
    "blue": "1F4E79",
    "teal": "0F766E",
    "gold": "B7791F",
    "red": "B42318",
    "light_blue": "EAF3FF",
    "light_teal": "E7F7F4",
    "light_gold": "FFF5D6",
    "light_gray": "F5F7FA",
    "border": "D6DAE3",
}


@dataclass
class Lesson:
    number: int
    start_page: int
    end_page: int
    raw: str
    terms: str
    patterns: str
    conversation: str
    grammar: str


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = COLORS["border"], size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_paragraph_bidi(paragraph, rtl: bool = True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if rtl and bidi is None:
        p_pr.append(OxmlElement("w:bidi"))
    elif not rtl and bidi is not None:
        p_pr.remove(bidi)


def set_run_font(run, ascii_font="Arial", east_asia_font="Yu Gothic") -> None:
    run.font.name = ascii_font
    r_pr = run._element.rPr
    if r_pr is None:
        r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:cs"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def add_text(paragraph, text: str, *, bold=False, size=11, color=COLORS["ink"], font="Arial"):
    run = paragraph.add_run(text)
    set_run_font(run, font)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def clean_text(text: str) -> str:
    text = text.replace("\u06be", "ه")
    text = re.sub(r"(?m)^\s*\d{1,3}\s*$", "", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def compact_lines(text: str, max_lines: int | None = None) -> list[str]:
    text = clean_text(text)
    lines: list[str] = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if len(line) > 150:
            parts = re.split(r"(?<=[.؟。])\s+", line)
            if len(parts) > 1:
                lines.extend(p.strip() for p in parts if p.strip())
            else:
                lines.append(line)
        else:
            lines.append(line)
    return lines[:max_lines] if max_lines else lines


def first_index(text: str, needles: list[str], start: int = 0) -> int:
    found = [idx for n in needles if (idx := text.find(n, start)) != -1]
    return min(found) if found else -1


def split_pages(extract: str) -> list[tuple[int, str]]:
    pages = []
    for chunk in extract.split("[[PAGE ")[1:]:
        page_no = int(chunk.split("]]", 1)[0])
        pages.append((page_no, chunk.split("]]", 1)[1]))
    return pages


def lesson_starts(pages: list[tuple[int, str]]) -> list[tuple[int, int]]:
    starts = []
    for page_no, content in pages:
        if LESSON in content and TERMS in content:
            starts.append((page_no, content.find(LESSON)))
    return starts


def build_lessons(pages: list[tuple[int, str]], starts: list[tuple[int, int]]) -> list[Lesson]:
    by_page = {page: text for page, text in pages}
    lessons = []
    for idx, (start_page, _) in enumerate(starts):
        end_page = (starts[idx + 1][0] - 1) if idx + 1 < len(starts) else 179
        raw_parts = []
        for page in range(start_page, end_page + 1):
            text = by_page.get(page, "")
            if page == start_page:
                lesson_idx = text.find(LESSON)
                text = text[lesson_idx:] if lesson_idx != -1 else text
            raw_parts.append(text)
        raw = clean_text("\n".join(raw_parts))

        terms_i = first_index(raw, [TERMS])
        patterns_i = first_index(raw, [TRANSLATION, PATTERNS], terms_i + 1)
        conv_i = first_index(raw, [CONVERSATION], patterns_i + 1 if patterns_i != -1 else 0)
        grammar_i = first_index(raw, [GRAMMAR], conv_i + 1 if conv_i != -1 else 0)

        terms_text = raw[terms_i:patterns_i] if terms_i != -1 and patterns_i != -1 else raw[:1800]
        patterns_text = raw[patterns_i:conv_i] if patterns_i != -1 and conv_i != -1 else ""
        conv_text = raw[conv_i:grammar_i] if conv_i != -1 and grammar_i != -1 else ""
        grammar_text = raw[grammar_i:] if grammar_i != -1 else raw[-2200:]

        lessons.append(
            Lesson(
                number=idx + 1,
                start_page=start_page,
                end_page=end_page,
                raw=raw,
                terms=terms_text,
                patterns=patterns_text,
                conversation=conv_text,
                grammar=grammar_text,
            )
        )
    return lessons


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 18, COLORS["blue"], 18, 8),
        ("Heading 2", 14, COLORS["teal"], 12, 6),
        ("Heading 3", 11.5, COLORS["gold"], 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def paragraph(doc: Document, text: str = "", *, style=None, rtl=True, align=None):
    p = doc.add_paragraph(style=style)
    if rtl:
        set_paragraph_bidi(p, True)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if align is not None:
        p.alignment = align
    if text:
        add_text(p, text)
    return p


def heading(doc: Document, text: str, level: int = 1, *, page_break=False):
    if page_break:
        doc.add_page_break()
    p = paragraph(doc, rtl=True, style=f"Heading {level}")
    add_text(p, text, bold=True, size={1: 18, 2: 14, 3: 11.5}[level], color={1: COLORS["blue"], 2: COLORS["teal"], 3: COLORS["gold"]}[level])
    return p


def callout(doc: Document, title: str, body: str, fill: str = COLORS["light_blue"], limit: int | None = None) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, fill)
    set_cell_border(cell)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    set_paragraph_bidi(p, True)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(p, title, bold=True, size=11.5, color=COLORS["blue"])
    for line in compact_lines(body, limit):
        p = cell.add_paragraph()
        set_paragraph_bidi(p, True)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_text(p, line, size=9.6)
    doc.add_paragraph()


def add_lines_block(doc: Document, lines: list[str], *, columns=1, max_lines=80) -> None:
    lines = lines[:max_lines]
    if not lines:
        paragraph(doc, "لا يوجد نص واضح مستخرج من المصدر لهذا القسم.")
        return
    if columns == 1:
        for line in lines:
            p = paragraph(doc, rtl=True)
            add_text(p, line, size=9.4)
        return
    rows = (len(lines) + columns - 1) // columns
    table = doc.add_table(rows=rows, cols=columns)
    table.style = "Table Grid"
    for r in range(rows):
        for c in range(columns):
            i = r + c * rows
            cell = table.cell(r, c)
            set_cell_border(cell)
            if i < len(lines):
                p = cell.paragraphs[0]
                set_paragraph_bidi(p, True)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                add_text(p, lines[i], size=8.8)
    doc.add_paragraph()


def japanese_terms(text: str, limit=8) -> list[str]:
    candidates = re.findall(r"[\u3040-\u30ff\u3400-\u9fff][\u3040-\u30ff\u3400-\u9fff\uff08\uff09\u30fc\u3001\u3002\s]{0,20}", text)
    cleaned = []
    for item in candidates:
        item = re.sub(r"\s+", " ", item).strip()
        if 1 <= len(item) <= 18 and item not in cleaned:
            cleaned.append(item)
    return cleaned[:limit]


def arabic_phrases(text: str, limit=8) -> list[str]:
    candidates = re.findall(r"[\u0621-\u064a][\u0621-\u064a\s،]{3,28}", text)
    cleaned = []
    for item in candidates:
        item = re.sub(r"\s+", " ", item).strip(" ،")
        if len(item) >= 4 and item not in cleaned and item not in [LESSON, TERMS]:
            cleaned.append(item)
    return cleaned[:limit]


def exercise_section(doc: Document, lesson: Lesson) -> None:
    heading(doc, "تمارين الدرس", 2)
    jp = japanese_terms(lesson.terms + "\n" + lesson.patterns, 10)
    ar = arabic_phrases(lesson.terms + "\n" + lesson.patterns, 10)
    particles = ["は", "か", "も", "これ", "それ", "あれ", "に", "へ", "で", "を", "が", "から", "まで", "より", "たい", "て", "たら"]

    exercises = [
        ("1. اختيار سريع", [
            f"اختر القراءة أو المعنى الأقرب لـ: {jp[0] if jp else '日本語'}",
            f"ما وظيفة النمط الأساسي في هذا الدرس؟ استخرجها من قسم القواعد.",
            f"اختر المفردة التي تناسب المعنى: {ar[0] if ar else 'كلمة من الدرس'}.",
        ]),
        ("2. املأ الفراغ", [
            f"{jp[1] if len(jp) > 1 else 'わたし'} ___ {jp[2] if len(jp) > 2 else '学生'} です。",
            f"{jp[3] if len(jp) > 3 else '日本'} ___ 行きます。",
            f"{jp[4] if len(jp) > 4 else '本'} ___ 読みます。",
        ]),
        ("3. ترجم إلى العربية", [
            jp[5] if len(jp) > 5 else "日本語を勉強します。",
            jp[6] if len(jp) > 6 else "これは本です。",
            jp[7] if len(jp) > 7 else "毎日練習します。",
        ]),
        ("4. كوّن جملة", [
            "استخدم ثلاث مفردات من الجدول واكتب جملة يابانية قصيرة.",
            "اكتب سؤالاً وجواباً باستعمال القاعدة الأساسية في الدرس.",
            "حوّل مثالاً واحداً من قسم الترجمة إلى جملة عن نفسك.",
        ]),
        ("5. مراجعة القاعدة", [
            f"استعمل إحدى هذه الأدوات في جملة صحيحة: {' / '.join(particles[(lesson.number - 1) % 8:(lesson.number - 1) % 8 + 4])}",
            "اشرح بالعربية متى تستخدم القاعدة الأهم في الدرس بجملة واحدة.",
        ]),
    ]

    for title, qs in exercises:
        p = paragraph(doc, rtl=True)
        add_text(p, title, bold=True, color=COLORS["gold"], size=10.8)
        for q in qs:
            p = paragraph(doc, rtl=True)
            add_text(p, f"□ {q}", size=9.8)
            p = paragraph(doc, rtl=True)
            add_text(p, "الإجابة: ................................................................................................", size=9, color=COLORS["muted"])

    callout(
        doc,
        "مفتاح إجابة مقترح",
        "تعتمد إجابات الترجمة وتكوين الجمل على مفردات الدرس وقواعده. راجع قسم المصطلحات والقواعد أعلاه، وتأكد من ترتيب الجملة اليابانية: الموضوع، الزمن/المكان، المفعول، ثم الفعل في النهاية عند الحاجة.",
        fill=COLORS["light_gold"],
        limit=4,
    )


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = paragraph(doc, rtl=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "اليابانية من طوكيو", bold=True, size=30, color=COLORS["blue"])
    p = paragraph(doc, rtl=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "日本語 From Tokyo", bold=True, size=24, color=COLORS["teal"])
    p = paragraph(doc, rtl=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "نسخة Word محسنة ومنظمة للدروس الـ 25", size=14, color=COLORS["muted"])
    for _ in range(2):
        doc.add_paragraph()
    callout(
        doc,
        "منهج الكتاب",
        "تم الحفاظ على جوهر كل درس: المصطلحات، أنماط الجملة، المحادثة، والقواعد. أضيفت تمارين تطبيقية بعد كل درس لتثبيت المفردات والقواعد عملياً.",
        fill=COLORS["light_teal"],
        limit=5,
    )
    doc.add_page_break()


def add_front_matter(doc: Document, pages: list[tuple[int, str]]) -> None:
    heading(doc, "تمهيد واستخدام الكتاب", 1)
    intro = "\n".join(text for page, text in pages if 3 <= page <= 10)
    callout(doc, "ملاحظة تحريرية", "هذه النسخة تعيد ترتيب المادة بصياغة نظيفة وتباعد واضح، مع حذف التراكب البصري والزخرفة الثقيلة حتى يكون الكتاب مناسباً للقراءة والطباعة.", fill=COLORS["light_blue"], limit=6)
    add_lines_block(doc, compact_lines(intro, 90), max_lines=90)
    doc.add_page_break()

    heading(doc, "أساسيات قبل الدروس", 1)
    basics = "\n".join(text for page, text in pages if 11 <= page <= 26)
    callout(doc, "مرجع سريع", "هذا القسم يجمع الشخصيات، الأدوات، ونظام الكتابة والتمهيد العام قبل الدخول إلى الدروس.", fill=COLORS["light_teal"], limit=4)
    add_lines_block(doc, compact_lines(basics, 150), columns=2, max_lines=150)


def add_lesson(doc: Document, lesson: Lesson) -> None:
    heading(doc, f"الدرس {lesson.number}", 1, page_break=True)
    p = paragraph(doc, rtl=True)
    add_text(p, f"صفحات المصدر: {lesson.start_page}-{lesson.end_page}", color=COLORS["muted"], size=9)

    callout(
        doc,
        "هدف الدرس",
        "اقرأ المصطلحات أولاً، ثم راجع أنماط الجملة، وبعدها طبّق القواعد في المحادثة والتمارين. ركّز على استعمال المفردات داخل جملة، وليس حفظها منفردة فقط.",
        fill=COLORS["light_blue"],
        limit=5,
    )

    heading(doc, "المفردات والمصطلحات", 2)
    add_lines_block(doc, compact_lines(lesson.terms, 95), columns=2, max_lines=95)

    heading(doc, "أنماط الجملة والترجمة", 2)
    add_lines_block(doc, compact_lines(lesson.patterns, 70), max_lines=70)

    heading(doc, "المحادثة", 2)
    add_lines_block(doc, compact_lines(lesson.conversation, 55), max_lines=55)

    heading(doc, "القواعد والشرح", 2)
    add_lines_block(doc, compact_lines(lesson.grammar, 110), max_lines=110)

    exercise_section(doc, lesson)


def add_appendix(doc: Document, pages: list[tuple[int, str]]) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    heading(doc, "ملحق الكانجي والمراجعة", 1)
    appendix = "\n".join(text for page, text in pages if 174 <= page <= 181)
    callout(doc, "طريقة الاستخدام", "استعمل الملحق كمراجعة بعد إنهاء الدروس. اكتب الكانجي والقراءات في دفتر منفصل ثم ارجع للأمثلة داخل الدروس.", fill=COLORS["light_gold"], limit=5)
    add_lines_block(doc, compact_lines(appendix, 180), columns=2, max_lines=180)


def main() -> None:
    extract = SOURCE_TEXT.read_text(encoding="utf-8")
    pages = split_pages(extract)
    starts = lesson_starts(pages)
    if len(starts) != 25:
        raise RuntimeError(f"Expected 25 lessons, found {len(starts)}")
    lessons = build_lessons(pages, starts)

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_front_matter(doc, pages)
    for lesson in lessons:
        add_lesson(doc, lesson)
    add_appendix(doc, pages)
    doc.save(OUTPUT)
    print(str(OUTPUT).encode("unicode_escape").decode("ascii"))


if __name__ == "__main__":
    main()
